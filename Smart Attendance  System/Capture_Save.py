import cv2
import docx
from docx.shared import Inches


cam = cv2.VideoCapture(0)
cv2.namedWindow("test")

img_counter = 0
while True:
    ret, frame = cam.read()
    if not ret:
        print("failed to grab frame")
        break
    cv2.imshow("test", frame)

    k = cv2.waitKey(1)
    if k%256 == 27:
        # ESC pressed
        print("Escape hit, closing...")
        break

    elif k%256 == 32:
        # SPACE pressed
        img_name = "opencv_frame_{}.png".format(img_counter)
        cv2.imwrite(img_name, frame)
        print("{} written!".format(img_name))
        img_counter += 1
     

        # SAVING IMAGE TO WORD FILE
        # Create an instance of a word document
        doc = docx.Document()

        # Add a Title to the document
        doc.add_heading('GeeksForGeeks', 0)


        # Image with defined size
        doc.add_heading('Student:', 3)
        doc.add_picture(img_name, width=Inches(2), height=Inches(2))

        # Now save the document to a location
        doc.save('Attendance_Img.docx')

        cam.release()
        cv2.destroyAllWindows()



