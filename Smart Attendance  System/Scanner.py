import cv2
import numpy as np
from pyzbar.pyzbar import decode
import sys
import csv


def decoder(image):
    gray_img = cv2.cvtColor(image,0)
    barcode = decode(gray_img)

    for obj in barcode:
        points = obj.polygon
        (x,y,w,h) = obj.rect
        pts = np.array(points, np.int32)
        pts = pts.reshape((-1, 1, 2))
        cv2.polylines(image, [pts], True, (0, 255, 0), 3)

        barcodeData = obj.data.decode("utf-8")
        barcodeType = obj.type
        string = "Data " + str(barcodeData) + " | Type " + str(barcodeType)
        
        cv2.putText(frame, string, (x,y), cv2.FONT_HERSHEY_SIMPLEX,0.8,(255,0,0), 2)
        print("Barcode: "+barcodeData +" | Type: "+barcodeType)

        #Storing Record in Docx File
        # Import docx NOT python-docx
        import docx

        # Create an instance of a word document
        doc = docx.Document()

        # Add a Title to the document
        doc.add_heading('Class Attendance', 0)

        # Image in its native size
        doc.add_heading('Student:', 3)
        doc.add_heading(barcodeData,3)

        # Now save the document to a location
        doc.save('Attendance.docx')

        #//////////////////////////////////////////////////////////////
        #Storing Record in Text File 
        # file = open('Attendance.txt', 'a+')
        # file.write("\n")
        # file.write(barcodeData)
        # file.close()

        #//////////////////////////////////////////////////////////////
        #Storing Record in CSV File 
        # # name of csv file
        # filename = "university_records.csv"
	
        # # writing to csv file
        # with open(filename, 'w') as csvfile:
        #     # creating a csv writer object
        #     csvwriter = csv.writer(csvfile)
                
        #     # writing the fields
        #     csvwriter.writerow(barcodeData)

        #//////////////////////////////////////////////////////////////




        # Terminating Window after one read
        sys.exit(0)
       


cap = cv2.VideoCapture(0)
while True:
    ret, frame = cap.read()
    decoder(frame)
    cv2.imshow('Image', frame)
    code = cv2.waitKey(10)
    if code == ord('q'):
        break




